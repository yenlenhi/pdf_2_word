"""
Focused PDF -> Word conversion pipeline.

Goals:
- Prefer direct PDF text extraction for speed and accuracy.
- Fall back to OCR only for pages without a usable text layer.
- Preserve document structure and tables at a practical level.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Iterable, List, Optional, Sequence, Tuple, Union

from PIL import Image

try:
    import cv2
    import numpy as np

    CV_AVAILABLE = True
except ImportError:
    cv2 = None
    np = None
    CV_AVAILABLE = False

try:
    import fitz

    PYMUPDF_AVAILABLE = True
except ImportError:
    fitz = None
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Cm, Pt

    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    WD_TABLE_ALIGNMENT = None
    WD_PARAGRAPH_ALIGNMENT = None
    Cm = None
    Pt = None
    DOCX_AVAILABLE = False


PdfInput = Union[str, Path, bytes, io.BytesIO]
BBox = Tuple[float, float, float, float]
VIETNAMESE_CHARS = set(
    "àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩị"
    "òóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ"
    "ÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊ"
    "ÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴĐ"
)
COMMON_VIETNAMESE_WORDS = {
    "và",
    "của",
    "cho",
    "với",
    "được",
    "là",
    "có",
    "không",
    "này",
    "đó",
    "một",
    "những",
    "như",
    "khi",
    "để",
    "từ",
    "trong",
    "nếu",
    "thì",
    "nhưng",
    "cũng",
    "còn",
    "rất",
    "đã",
    "sẽ",
    "ngày",
    "năm",
    "thông",
    "báo",
    "tuyển",
    "sinh",
    "đào",
    "tạo",
    "trường",
    "đại",
    "học",
    "công",
    "an",
    "chính",
    "quy",
}


@dataclass
class TextBlock:
    bbox: BBox
    text: str
    font_size: float = 11.0
    is_bold: bool = False
    kind: str = "text"
    line_lefts: List[float] = field(default_factory=list)
    line_rights: List[float] = field(default_factory=list)


@dataclass
class TableBlock:
    bbox: BBox
    rows: List[List[str]]
    source: str = "pymupdf"


@dataclass
class PageContent:
    page_num: int
    width: float
    height: float
    text_blocks: List[TextBlock] = field(default_factory=list)
    tables: List[TableBlock] = field(default_factory=list)
    used_ocr: bool = False
    strategy: str = "direct_text"


@dataclass
class PDFToWordConversionResult:
    success: bool
    docx_bytes: bytes = b""
    extracted_text: str = ""
    pages_converted: int = 0
    table_count: int = 0
    ocr_pages: int = 0
    error_message: str = ""
    method_used: str = ""
    pages: List[PageContent] = field(default_factory=list)


class PDFToWordConverter:
    DOC_PAGE_WIDTH_CM = 21.0
    DOC_PAGE_HEIGHT_CM = 29.7
    DOC_MARGIN_TOP_CM = 1.8
    DOC_MARGIN_BOTTOM_CM = 1.8
    DOC_MARGIN_LEFT_CM = 2.0
    DOC_MARGIN_RIGHT_CM = 2.0
    POINTS_PER_CM = 28.3464567
    SENTENCE_ENDINGS = (".", ":", ";", "?", "!")

    def __init__(self, ocr_callback: Optional[Callable[[Image.Image], str]] = None):
        self.ocr_callback = ocr_callback

    def convert(
        self,
        pdf_input: PdfInput,
        output_path: Optional[str] = None,
        start_page: int = 0,
        end_page: Optional[int] = None,
        dpi: int = 220,
        extract_tables: bool = True,
        progress_callback: Optional[Callable[[float, str], None]] = None,
    ) -> PDFToWordConversionResult:
        if not PYMUPDF_AVAILABLE:
            return PDFToWordConversionResult(
                success=False,
                error_message="PyMuPDF is not installed.",
            )
        if not DOCX_AVAILABLE:
            return PDFToWordConversionResult(
                success=False,
                error_message="python-docx is not installed.",
            )

        try:
            doc = self._open_document(pdf_input)
        except Exception as exc:
            return PDFToWordConversionResult(success=False, error_message=str(exc))

        try:
            total_pages = len(doc)
            if total_pages == 0:
                return PDFToWordConversionResult(success=False, error_message="PDF has no pages.")

            end_page = total_pages if end_page is None else min(end_page, total_pages)
            start_page = max(0, start_page)
            if start_page >= end_page:
                return PDFToWordConversionResult(success=False, error_message="Invalid page range.")

            pages: List[PageContent] = []
            for page_index in range(start_page, end_page):
                page = doc.load_page(page_index)
                if progress_callback:
                    progress = (page_index - start_page) / max(1, end_page - start_page)
                    progress_callback(progress, f"Processing page {page_index + 1}/{end_page}...")
                pages.append(
                    self._extract_page_content(
                        page,
                        dpi=dpi,
                        extract_tables=extract_tables,
                    )
                )

            docx = self._build_document(pages)
            docx_bytes = self._document_to_bytes(docx)
            extracted_text = self._pages_to_text(pages)

            if output_path:
                with open(output_path, "wb") as file_obj:
                    file_obj.write(docx_bytes)

            ocr_pages = sum(1 for page in pages if page.used_ocr)
            table_count = sum(len(page.tables) for page in pages)
            method_used = "direct_text"
            if ocr_pages:
                method_used = "hybrid"

            if progress_callback:
                progress_callback(1.0, "Word document ready.")

            return PDFToWordConversionResult(
                success=True,
                docx_bytes=docx_bytes,
                extracted_text=extracted_text,
                pages_converted=len(pages),
                table_count=table_count,
                ocr_pages=ocr_pages,
                method_used=method_used,
                pages=pages,
            )
        except Exception as exc:
            return PDFToWordConversionResult(success=False, error_message=str(exc))
        finally:
            doc.close()

    def _open_document(self, pdf_input: PdfInput):
        if isinstance(pdf_input, (str, Path)):
            return fitz.open(str(pdf_input))
        if isinstance(pdf_input, bytes):
            return fitz.open(stream=pdf_input, filetype="pdf")
        if isinstance(pdf_input, io.BytesIO):
            pdf_input.seek(0)
            return fitz.open(stream=pdf_input.read(), filetype="pdf")
        raise TypeError(f"Unsupported PDF input type: {type(pdf_input)}")

    def _extract_page_content(self, page, dpi: int, extract_tables: bool) -> PageContent:
        page_content = PageContent(
            page_num=page.number + 1,
            width=float(page.rect.width),
            height=float(page.rect.height),
        )

        content = page.get_text("dict", sort=True)
        tables = self._extract_tables(page) if extract_tables else []
        text_blocks = self._extract_text_blocks(content, tables, page_content.width)
        image_blocks = self._extract_image_ocr_blocks(content)
        strategy = self._classify_page(text_blocks, image_blocks)
        if self._should_force_ocr(text_blocks):
            strategy = "garbled_text_ocr"

        if strategy in {"scanned_ocr", "garbled_text_ocr"}:
            ocr_blocks = self._ocr_page_blocks(page, dpi=dpi)
            if ocr_blocks:
                page_content.used_ocr = True
                page_content.strategy = strategy
                page_content.text_blocks = ocr_blocks
                page_content.tables = [] if strategy == "scanned_ocr" else tables
                return page_content

        text_blocks.extend(image_blocks)
        text_blocks.sort(key=lambda block: (block.bbox[1], block.bbox[0]))

        if strategy == "ocr_fallback":
            ocr_blocks = self._ocr_page_blocks(page, dpi=dpi)
            if ocr_blocks:
                page_content.used_ocr = True
                page_content.strategy = strategy
                page_content.text_blocks = ocr_blocks
                page_content.tables = tables
                return page_content

        page_content.strategy = strategy
        page_content.text_blocks = text_blocks
        page_content.tables = tables
        return page_content

    def _extract_text_blocks(
        self,
        content: dict,
        tables: Sequence[TableBlock],
        page_width: float,
    ) -> List[TextBlock]:
        line_blocks: List[TextBlock] = []

        for block in content.get("blocks", []):
            if block.get("type") != 0:
                continue

            for line in block.get("lines", []):
                line_bbox_raw = line.get("bbox") or block.get("bbox", (0, 0, 0, 0))
                line_bbox = tuple(float(value) for value in line_bbox_raw[:4])
                if any(self._overlap_ratio(line_bbox, table.bbox) > 0.55 for table in tables):
                    continue

                spans = line.get("spans", [])
                line_text = self._join_line_spans(spans)
                if not line_text:
                    continue

                font_sizes: List[float] = []
                bold_hits = 0
                for span in spans:
                    size = span.get("size")
                    if size:
                        font_sizes.append(float(size))
                    if "bold" in str(span.get("font", "")).lower():
                        bold_hits += 1

                font_size = max(font_sizes) if font_sizes else 11.0
                is_bold = bold_hits > 0 and bold_hits >= max(1, len(spans) // 2)
                line_blocks.append(
                TextBlock(
                        bbox=line_bbox,
                        text=line_text,
                        font_size=font_size,
                        is_bold=is_bold,
                        kind="text",
                        line_lefts=[line_bbox[0]],
                        line_rights=[line_bbox[2]],
                    )
                )

        return self._merge_direct_text_lines(line_blocks, page_width)

    def _join_line_spans(self, spans: Sequence[dict]) -> str:
        pieces: List[str] = []
        prev_right: Optional[float] = None
        prev_text = ""

        for span in spans:
            text = str(span.get("text", ""))
            if not text:
                continue

            bbox = span.get("bbox") or (0, 0, 0, 0)
            left = float(bbox[0]) if len(bbox) >= 1 else 0.0
            right = float(bbox[2]) if len(bbox) >= 3 else left
            font_size = float(span.get("size") or 11.0)

            if pieces and prev_right is not None:
                gap = left - prev_right
                needs_space = (
                    gap > max(1.5, font_size * 0.18)
                    and prev_text
                    and text
                    and prev_text[-1].isalnum()
                    and text[0].isalnum()
                )
                if needs_space:
                    pieces.append(" ")

            pieces.append(text)
            prev_right = right
            prev_text = text

        return self._normalize_text("".join(pieces))

    def _merge_direct_text_lines(
        self,
        line_blocks: Sequence[TextBlock],
        page_width: float,
    ) -> List[TextBlock]:
        if not line_blocks:
            return []

        ordered = sorted(line_blocks, key=lambda block: (block.bbox[1], block.bbox[0]))
        merged: List[TextBlock] = [self._clone_block(ordered[0])]

        for line in ordered[1:]:
            previous = merged[-1]
            if self._should_merge_direct_line(previous, line, page_width):
                previous.text = f"{previous.text}\n{line.text}"
                previous.bbox = (
                    min(previous.bbox[0], line.bbox[0]),
                    min(previous.bbox[1], line.bbox[1]),
                    max(previous.bbox[2], line.bbox[2]),
                    max(previous.bbox[3], line.bbox[3]),
                )
                previous.font_size = max(previous.font_size, line.font_size)
                previous.is_bold = previous.is_bold and line.is_bold
                previous.line_lefts.extend(line.line_lefts)
                previous.line_rights.extend(line.line_rights)
            else:
                merged.append(self._clone_block(line))

        for block in merged:
            block.kind = self._classify_text_block_kind(block, page_width)

        return merged

    def _clone_block(self, block: TextBlock) -> TextBlock:
        return TextBlock(
            bbox=block.bbox,
            text=block.text,
            font_size=block.font_size,
            is_bold=block.is_bold,
            kind=block.kind,
            line_lefts=list(block.line_lefts),
            line_rights=list(block.line_rights),
        )

    def _should_merge_direct_line(
        self,
        previous: TextBlock,
        current: TextBlock,
        page_width: float,
    ) -> bool:
        previous_lines = [line.strip() for line in previous.text.splitlines() if line.strip()]
        current_text = current.text.strip()
        if not previous_lines or not current_text:
            return False
        if self._is_list_block([previous_lines[-1]]) or self._is_list_block([current_text]):
            return False
        if self._is_heading_candidate(previous, page_width) or self._is_heading_candidate(current, page_width):
            return False

        vertical_gap = current.bbox[1] - previous.bbox[3]
        max_gap = max(5.0, min(previous.font_size, current.font_size) * 0.9)
        if vertical_gap > max_gap:
            return False

        left_tolerance = max(14.0, page_width * 0.025)
        previous_body_left = min(previous.line_lefts) if previous.line_lefts else previous.bbox[0]
        current_left = current.line_lefts[0] if current.line_lefts else current.bbox[0]
        if abs(current_left - previous_body_left) > left_tolerance:
            return False

        font_delta = abs(previous.font_size - current.font_size)
        if font_delta > 1.6:
            return False

        previous_last_line = previous_lines[-1]
        previous_width = (previous.line_rights[-1] - previous.line_lefts[-1]) if previous.line_lefts and previous.line_rights else (previous.bbox[2] - previous.bbox[0])
        width_ratio = previous_width / max(1.0, page_width)
        if previous_last_line.endswith(self.SENTENCE_ENDINGS) and width_ratio < 0.58:
            return False

        return True

    def _extract_image_ocr_blocks(self, content: dict) -> List[TextBlock]:
        if not self.ocr_callback:
            return []

        blocks: List[TextBlock] = []
        for block in content.get("blocks", []):
            if block.get("type") != 1:
                continue

            bbox = tuple(float(value) for value in block.get("bbox", (0, 0, 0, 0)))
            width = bbox[2] - bbox[0]
            height = bbox[3] - bbox[1]
            if width < 40 or height < 20:
                continue

            image_bytes = block.get("image")
            if not image_bytes:
                continue

            try:
                image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                text = self._normalize_text(self.ocr_callback(image) or "")
            except Exception:
                text = ""

            if not text:
                continue

            blocks.append(
                TextBlock(
                    bbox=bbox,
                    text=text,
                    font_size=11.0,
                    is_bold=False,
                    kind="ocr",
                    line_lefts=[],
                    line_rights=[],
                )
            )

        return blocks

    def _classify_page(
        self,
        text_blocks: Sequence[TextBlock],
        image_blocks: Sequence[TextBlock],
    ) -> str:
        text_len = sum(len(block.text.strip()) for block in text_blocks)
        image_text_len = sum(len(block.text.strip()) for block in image_blocks)

        if text_len >= 120:
            return "mixed" if image_text_len >= 20 else "direct_text"
        if text_len >= 40:
            return "ocr_fallback" if self.ocr_callback else "direct_text"
        if image_text_len >= 80:
            return "mixed"
        return "scanned_ocr" if self.ocr_callback else "direct_text"

    def _should_force_ocr(self, text_blocks: Sequence[TextBlock]) -> bool:
        if not self.ocr_callback or not text_blocks:
            return False

        combined_text = self._normalize_text("\n".join(block.text for block in text_blocks if block.text))
        if len(combined_text) < 80:
            return False

        return self._text_quality_score(combined_text) < 0.45

    def _text_quality_score(self, text: str) -> float:
        text = self._normalize_text(text)
        if not text:
            return 0.0

        tokens = re.findall(r"\S+", text)
        if not tokens:
            return 0.0

        alpha_tokens = [
            token
            for token in tokens
            if re.search(r"[A-Za-zÀ-ỹĐđ]", token)
        ]
        if not alpha_tokens:
            return 0.2

        suspicious_tokens = 0
        long_upper_tokens = 0
        mixed_alnum_tokens = 0
        apostrophe_tokens = 0
        common_hits = 0

        for raw_token in alpha_tokens:
            token = raw_token.strip("()[]{}\"'“”‘’.,;:!?")
            lowered = token.lower()

            if lowered in COMMON_VIETNAMESE_WORDS:
                common_hits += 1

            if re.search(r"[A-Za-zÀ-ỹĐđ]\d|\d[A-Za-zÀ-ỹĐđ]", token):
                mixed_alnum_tokens += 1
                suspicious_tokens += 1

            if "'" in token or "’" in token:
                apostrophe_tokens += 1

            if len(token) >= 14 and re.fullmatch(r"[A-ZĐÀ-Ỵ]+", token):
                long_upper_tokens += 1
                suspicious_tokens += 1

            if len(token) >= 18 and token.isalpha():
                suspicious_tokens += 1

        letters = [char for char in text if char.isalpha()]
        viet_ratio = (
            sum(1 for char in letters if char in VIETNAMESE_CHARS) / max(1, len(letters))
        )
        suspicious_ratio = suspicious_tokens / max(1, len(alpha_tokens))
        mixed_ratio = mixed_alnum_tokens / max(1, len(alpha_tokens))
        long_upper_ratio = long_upper_tokens / max(1, len(alpha_tokens))
        apostrophe_ratio = apostrophe_tokens / max(1, len(alpha_tokens))
        common_ratio = common_hits / max(1, len(alpha_tokens))

        score = 1.0
        score -= suspicious_ratio * 1.2
        score -= mixed_ratio * 1.4
        score -= long_upper_ratio * 0.9
        score -= apostrophe_ratio * 0.35
        score += min(0.2, common_ratio * 0.35)
        score += min(0.15, viet_ratio * 0.5)
        return max(0.0, min(1.0, score))

    def _extract_tables(self, page) -> List[TableBlock]:
        if not hasattr(page, "find_tables"):
            return []

        extracted: List[TableBlock] = []
        for strategy in ("lines_strict", "text"):
            try:
                found = page.find_tables(strategy=strategy)
            except TypeError:
                found = page.find_tables()
            except Exception:
                continue

            for table in getattr(found, "tables", []):
                rows = self._normalize_table_rows(table.extract())
                if not rows:
                    continue
                extracted.append(
                    TableBlock(
                        bbox=tuple(float(value) for value in table.bbox),
                        rows=rows,
                        source=f"pymupdf:{strategy}",
                    )
                )

            if extracted:
                break

        return self._dedupe_tables(extracted)

    def _normalize_table_rows(self, rows: Optional[Iterable[Iterable[object]]]) -> List[List[str]]:
        if not rows:
            return []

        normalized = [[self._clean_cell(cell) for cell in row] for row in rows]
        normalized = [row for row in normalized if any(cell for cell in row)]
        if not normalized:
            return []

        max_cols = max(len(row) for row in normalized)
        keep_cols = []
        for col_idx in range(max_cols):
            has_content = any(col_idx < len(row) and row[col_idx] for row in normalized)
            if has_content:
                keep_cols.append(col_idx)

        if not keep_cols:
            return []

        compact = []
        for row in normalized:
            compact.append([row[idx] if idx < len(row) else "" for idx in keep_cols])

        return compact

    def _clean_cell(self, value: object) -> str:
        if value is None:
            return ""
        text = str(value).replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _dedupe_tables(self, tables: Sequence[TableBlock]) -> List[TableBlock]:
        deduped: List[TableBlock] = []
        seen = set()
        for table in tables:
            key = (
                round(table.bbox[0], 1),
                round(table.bbox[1], 1),
                round(table.bbox[2], 1),
                round(table.bbox[3], 1),
                len(table.rows),
                max((len(row) for row in table.rows), default=0),
            )
            if key in seen:
                continue
            seen.add(key)
            deduped.append(table)
        return deduped

    def _render_page_image(self, page, dpi: int) -> Tuple[Image.Image, float]:
        zoom = dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        return image, zoom

    def _ocr_page_blocks(self, page, dpi: int) -> List[TextBlock]:
        if not self.ocr_callback:
            return []

        image, zoom = self._render_page_image(page, dpi=dpi)
        regions = self._detect_text_regions(image)
        if not regions:
            text = self._normalize_text(self.ocr_callback(image) or "")
            if not text:
                return []
            return [
                TextBlock(
                    bbox=(0.0, 0.0, float(page.rect.width), float(page.rect.height)),
                    text=text,
                    font_size=11.0,
                    is_bold=False,
                    kind="ocr",
                    line_lefts=[],
                    line_rights=[],
                )
            ]

        blocks: List[TextBlock] = []
        for x1, y1, x2, y2 in self._order_regions(regions, image.width):
            crop = image.crop((x1, y1, x2, y2))
            text = self._normalize_text(self.ocr_callback(crop) or "")
            if not text:
                continue

            bbox = (
                x1 / zoom,
                y1 / zoom,
                x2 / zoom,
                y2 / zoom,
            )
            block_height = max(1.0, (y2 - y1) / zoom)
            blocks.append(
                TextBlock(
                    bbox=bbox,
                    text=text,
                    font_size=max(11.0, min(16.0, block_height * 0.55)),
                    is_bold=False,
                    kind="ocr",
                    line_lefts=[],
                    line_rights=[],
                )
            )

        merged = self._merge_ocr_blocks(blocks)
        return merged if merged else blocks

    def _detect_text_regions(self, image: Image.Image) -> List[Tuple[int, int, int, int]]:
        if not CV_AVAILABLE:
            return []

        rgb = np.array(image)
        gray = cv2.cvtColor(rgb, cv2.COLOR_RGB2GRAY)
        gray = cv2.GaussianBlur(gray, (3, 3), 0)
        binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

        h, w = binary.shape
        kernel_width = max(24, w // 35)
        kernel_height = max(3, h // 180)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (kernel_width, kernel_height))
        merged = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel, iterations=2)

        contours, _ = cv2.findContours(merged, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        regions: List[Tuple[int, int, int, int]] = []
        min_area = max(1200, int(w * h * 0.00012))
        for contour in contours:
            x, y, cw, ch = cv2.boundingRect(contour)
            area = cw * ch
            if area < min_area:
                continue
            if cw < 30 or ch < 12:
                continue
            if cw > w * 0.96 and ch > h * 0.9:
                continue
            regions.append((x, y, x + cw, y + ch))

        return regions

    def _order_regions(
        self, regions: Sequence[Tuple[int, int, int, int]], page_width: int
    ) -> List[Tuple[int, int, int, int]]:
        if not regions:
            return []

        left = [region for region in regions if (region[0] + region[2]) / 2 < page_width * 0.48]
        right = [region for region in regions if (region[0] + region[2]) / 2 >= page_width * 0.52]
        avg_width = sum(region[2] - region[0] for region in regions) / len(regions)

        if len(left) >= 3 and len(right) >= 3 and avg_width < page_width * 0.42:
            return sorted(left, key=lambda region: (region[1], region[0])) + sorted(
                right, key=lambda region: (region[1], region[0])
            )

        return sorted(regions, key=lambda region: (round(region[1] / 12), region[0]))

    def _merge_ocr_blocks(self, blocks: Sequence[TextBlock]) -> List[TextBlock]:
        if not blocks:
            return []

        merged: List[TextBlock] = [blocks[0]]
        for block in blocks[1:]:
            prev = merged[-1]
            same_column = abs(prev.bbox[0] - block.bbox[0]) <= 24
            small_gap = 0 <= block.bbox[1] - prev.bbox[3] <= 10
            similar_width = abs((prev.bbox[2] - prev.bbox[0]) - (block.bbox[2] - block.bbox[0])) <= 80

            if same_column and small_gap and similar_width:
                prev.text = f"{prev.text}\n{block.text}"
                prev.bbox = (
                    min(prev.bbox[0], block.bbox[0]),
                    min(prev.bbox[1], block.bbox[1]),
                    max(prev.bbox[2], block.bbox[2]),
                    max(prev.bbox[3], block.bbox[3]),
                )
                prev.font_size = max(prev.font_size, block.font_size)
                prev.line_lefts.extend(block.line_lefts)
                prev.line_rights.extend(block.line_rights)
            else:
                merged.append(block)

        return merged

    def _build_document(self, pages: Sequence[PageContent]):
        doc = Document()
        self._configure_document(doc)

        for index, page in enumerate(pages):
            blocks = []
            for text_block in page.text_blocks:
                blocks.append((text_block.bbox[1], "text", text_block))
            for table_block in page.tables:
                blocks.append((table_block.bbox[1], "table", table_block))
            blocks.sort(key=lambda item: item[0])
            previous_bottom: Optional[float] = None
            page_left_ref, page_right_ref = self._estimate_page_text_bounds(page)

            if not blocks:
                doc.add_paragraph("")

            for _, block_type, block in blocks:
                if block_type == "text":
                    self._write_text_block(
                        doc,
                        block,
                        page.width,
                        page.height,
                        page_left_ref,
                        page_right_ref,
                        previous_bottom=previous_bottom,
                    )
                else:
                    self._write_table(doc, block)
                previous_bottom = block.bbox[3]

            if index < len(pages) - 1:
                doc.add_page_break()

        return doc

    def _configure_document(self, doc) -> None:
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        for section in doc.sections:
            section.page_width = Cm(self.DOC_PAGE_WIDTH_CM)
            section.page_height = Cm(self.DOC_PAGE_HEIGHT_CM)
            section.top_margin = Cm(self.DOC_MARGIN_TOP_CM)
            section.bottom_margin = Cm(self.DOC_MARGIN_BOTTOM_CM)
            section.left_margin = Cm(self.DOC_MARGIN_LEFT_CM)
            section.right_margin = Cm(self.DOC_MARGIN_RIGHT_CM)

    def _write_text_block(
        self,
        doc,
        block: TextBlock,
        page_width: float,
        page_height: float,
        page_left_ref: float,
        page_right_ref: float,
        previous_bottom: Optional[float] = None,
    ) -> None:
        paragraph = doc.add_paragraph()
        self._apply_paragraph_geometry(
            paragraph,
            block,
            page_width,
            page_height,
            page_left_ref,
            page_right_ref,
            previous_bottom,
        )
        paragraph.alignment = self._get_paragraph_alignment(block, page_width)

        lines = self._prepare_output_lines(block)
        for line_index, line in enumerate(lines):
            run = paragraph.add_run(line)
            run.bold = block.is_bold or block.kind == "title"
            run.font.size = Pt(max(10, min(18, block.font_size)))
            if line_index < len(lines) - 1:
                run.add_break()

    def _write_table(self, doc, table: TableBlock) -> None:
        if not table.rows:
            return

        column_count = max(len(row) for row in table.rows)
        word_table = doc.add_table(rows=len(table.rows), cols=column_count)
        word_table.style = "Table Grid"
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_index, row in enumerate(table.rows):
            for col_index in range(column_count):
                value = row[col_index] if col_index < len(row) else ""
                cell = word_table.cell(row_index, col_index)
                cell.text = value
                if row_index == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        doc.add_paragraph("")

    def _document_to_bytes(self, doc) -> bytes:
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _apply_paragraph_geometry(
        self,
        paragraph,
        block: TextBlock,
        page_width: float,
        page_height: float,
        page_left_ref: float,
        page_right_ref: float,
        previous_bottom: Optional[float],
    ) -> None:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(4)

        lines = [line.strip() for line in block.text.splitlines() if line.strip()]
        should_center = self._should_center_block(block, page_width, lines)
        if should_center:
            paragraph_format.left_indent = Pt(0)
            paragraph_format.right_indent = Pt(0)
        else:
            left_indent_pt = self._scale_horizontal(max(0.0, block.bbox[0] - page_left_ref), page_width)
            right_indent_pt = self._scale_horizontal(
                max(0.0, page_right_ref - block.bbox[2]),
                page_width,
            )
            max_side_indent = self._content_width_pt() * 0.35
            paragraph_format.left_indent = Pt(min(left_indent_pt, max_side_indent))
            paragraph_format.right_indent = Pt(min(right_indent_pt, max_side_indent))

        paragraph_format.first_line_indent = Pt(self._estimate_first_line_indent(block, page_width))

        if previous_bottom is None:
            paragraph_format.space_before = Pt(0)
        else:
            gap = max(0.0, block.bbox[1] - previous_bottom)
            gap_pt = self._scale_vertical(gap, page_height)
            paragraph_format.space_before = Pt(min(18.0, gap_pt * 0.85))

    def _scale_horizontal(self, value: float, page_width: float) -> float:
        if page_width <= 0:
            return 0.0
        return max(0.0, value / page_width) * self._content_width_pt()

    def _scale_vertical(self, value: float, page_height: float) -> float:
        if page_height <= 0:
            return 0.0
        return max(0.0, value / page_height) * self._content_height_pt()

    def _content_width_pt(self) -> float:
        return (
            self.DOC_PAGE_WIDTH_CM - self.DOC_MARGIN_LEFT_CM - self.DOC_MARGIN_RIGHT_CM
        ) * self.POINTS_PER_CM

    def _content_height_pt(self) -> float:
        return (
            self.DOC_PAGE_HEIGHT_CM - self.DOC_MARGIN_TOP_CM - self.DOC_MARGIN_BOTTOM_CM
        ) * self.POINTS_PER_CM

    def _estimate_page_text_bounds(self, page: PageContent) -> Tuple[float, float]:
        candidates = [
            block
            for block in page.text_blocks
            if block.kind != "title" and (block.bbox[2] - block.bbox[0]) >= page.width * 0.35
        ]
        if not candidates:
            candidates = list(page.text_blocks)
        if not candidates:
            return 0.0, page.width

        left = min(block.bbox[0] for block in candidates)
        right = max(block.bbox[2] for block in candidates)
        if right <= left:
            return 0.0, page.width
        return left, right

    def _estimate_first_line_indent(self, block: TextBlock, page_width: float) -> float:
        if block.kind == "title":
            return 0.0

        lines = [line.strip() for line in block.text.splitlines() if line.strip()]
        if self._is_list_block(lines):
            return 0.0
        if len(block.line_lefts) < 2:
            return 0.0

        first_left = block.line_lefts[0]
        body_left = min(block.line_lefts[1:])
        delta = first_left - body_left
        if abs(delta) < page_width * 0.01:
            return 0.0

        max_indent_pt = self._content_width_pt() * 0.12
        scaled = self._scale_horizontal(delta, page_width)
        return max(-max_indent_pt, min(max_indent_pt, scaled))

    def _classify_text_block_kind(self, block: TextBlock, page_width: float) -> str:
        lines = [line.strip() for line in block.text.splitlines() if line.strip()]
        if self._is_list_block(lines):
            return "list"
        if self._is_heading_candidate(block, page_width):
            return "title"
        return "text"

    def _is_heading_candidate(self, block: TextBlock, page_width: float) -> bool:
        text = block.text.strip()
        if not text:
            return False
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines or self._is_list_block(lines):
            return False
        if self._ends_like_heading(text):
            return False

        width = max(1.0, block.bbox[2] - block.bbox[0])
        is_centered = self._should_center_block(block, page_width, lines)
        all_caps = self._is_mostly_upper(text)
        numbered_heading = bool(re.match(r"^([A-ZÄ]\.|[IVXLCDM]+\.)\s+|^\d+([.)]\d*)*\s+", lines[0]))

        if block.font_size >= 15 and len(text) <= 220:
            return True
        if is_centered and (all_caps or block.font_size >= 13):
            return True
        if numbered_heading and (block.is_bold or block.font_size >= 12):
            return True
        return False

    def _is_mostly_upper(self, text: str) -> bool:
        letters = [char for char in text if char.isalpha()]
        if not letters:
            return False
        uppercase_letters = [char for char in letters if char.isupper()]
        return len(uppercase_letters) / len(letters) >= 0.72

    def _prepare_output_lines(self, block: TextBlock) -> List[str]:
        raw_lines = [line.strip() for line in block.text.splitlines() if line.strip()]
        if not raw_lines:
            return [block.text.strip()] if block.text.strip() else [""]
        if block.kind == "title" or self._is_list_block(raw_lines):
            return raw_lines
        return [self._join_wrapped_lines(raw_lines)]

    def _join_wrapped_lines(self, lines: Sequence[str]) -> str:
        merged_parts: List[str] = []
        for line in lines:
            if not merged_parts:
                merged_parts.append(line)
                continue

            previous = merged_parts[-1]
            if previous.endswith("-"):
                merged_parts[-1] = previous[:-1] + line
            else:
                merged_parts[-1] = f"{previous} {line}"

        return " ".join(part.strip() for part in merged_parts if part.strip())

    def _get_paragraph_alignment(self, block: TextBlock, page_width: float):
        text = block.text.strip()
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not text:
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        if self._should_center_block(block, page_width, lines):
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        if self._is_list_block(lines):
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        if len(text) >= 60 and not self._ends_like_heading(text):
            return WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        return WD_PARAGRAPH_ALIGNMENT.LEFT

    def _should_center_block(
        self,
        block: TextBlock,
        page_width: float,
        lines: Sequence[str],
    ) -> bool:
        if block.kind != "title" or not lines:
            return False
        if self._is_list_block(lines):
            return False

        text = block.text.strip()
        if len(lines) > 2 or len(text) > 140:
            return False
        if self._ends_like_heading(text):
            return False

        block_left, _, block_right, _ = block.bbox
        block_width = max(1.0, block_right - block_left)
        block_center = (block_left + block_right) / 2
        center_offset = abs(block_center - (page_width / 2))
        left_margin = block_left
        right_margin = max(0.0, page_width - block_right)

        return (
            block_width <= page_width * 0.72
            and center_offset <= page_width * 0.08
            and abs(left_margin - right_margin) <= page_width * 0.12
        )

    def _is_list_block(self, lines: Sequence[str]) -> bool:
        return any(re.match(r"^([•\-*o○]|[0-9]+[.)]|[A-Za-z][.)])\s+", line) for line in lines)

    def _ends_like_heading(self, text: str) -> bool:
        return text.endswith((".", ":", ";", ","))

    def _pages_to_text(self, pages: Sequence[PageContent]) -> str:
        chunks = []
        for page in pages:
            page_parts = []
            combined = []
            for text_block in page.text_blocks:
                combined.append((text_block.bbox[1], "text", text_block.text))
            for table in page.tables:
                combined.append((table.bbox[1], "table", self._table_to_tsv(table.rows)))
            combined.sort(key=lambda item: item[0])
            for _, _, value in combined:
                if value.strip():
                    page_parts.append(value.strip())
            chunks.append("\n\n".join(page_parts).strip())
        return "\n\n".join(chunk for chunk in chunks if chunk).strip()

    def _table_to_tsv(self, rows: Sequence[Sequence[str]]) -> str:
        return "\n".join("\t".join(cell for cell in row) for row in rows)

    def _normalize_text(self, text: str) -> str:
        text = text.replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _looks_like_title(self, text: str, font_size: float, is_bold: bool) -> bool:
        clean_text = text.strip()
        if not clean_text:
            return False
        if re.match(r"^([•\-*o○]|[0-9]+[.)]|[A-Za-z][.)])\s+", clean_text):
            return False
        if clean_text.endswith((".", ":", ";", ",")):
            return False
        if font_size >= 15:
            return True
        if self._is_mostly_upper(clean_text) and len(clean_text) <= 180:
            return True
        if is_bold and len(clean_text) <= 80 and re.match(r"^\d+([.)]\d*)*\s+", clean_text):
            return True
        return False

    def _overlap_ratio(self, bbox_a: BBox, bbox_b: BBox) -> float:
        x1 = max(bbox_a[0], bbox_b[0])
        y1 = max(bbox_a[1], bbox_b[1])
        x2 = min(bbox_a[2], bbox_b[2])
        y2 = min(bbox_a[3], bbox_b[3])
        if x2 <= x1 or y2 <= y1:
            return 0.0
        inter_area = (x2 - x1) * (y2 - y1)
        area_a = max(1.0, (bbox_a[2] - bbox_a[0]) * (bbox_a[3] - bbox_a[1]))
        return inter_area / area_a


def convert_pdf_to_word(
    pdf_input: PdfInput,
    output_path: Optional[str] = None,
    start_page: int = 0,
    end_page: Optional[int] = None,
    dpi: int = 220,
    extract_tables: bool = True,
    ocr_callback: Optional[Callable[[Image.Image], str]] = None,
    progress_callback: Optional[Callable[[float, str], None]] = None,
) -> PDFToWordConversionResult:
    converter = PDFToWordConverter(ocr_callback=ocr_callback)
    return converter.convert(
        pdf_input=pdf_input,
        output_path=output_path,
        start_page=start_page,
        end_page=end_page,
        dpi=dpi,
        extract_tables=extract_tables,
        progress_callback=progress_callback,
    )


__all__ = [
    "DOCX_AVAILABLE",
    "PYMUPDF_AVAILABLE",
    "PDFToWordConversionResult",
    "PDFToWordConverter",
    "PageContent",
    "TableBlock",
    "TextBlock",
    "convert_pdf_to_word",
]
