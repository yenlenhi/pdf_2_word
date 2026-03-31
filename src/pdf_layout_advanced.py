"""
PDF Layout Analysis với Table Detection - Advanced Version
Phân tích layout PDF, giữ nguyên vị trí text và bảng

Technologies:
- Classical Computer Vision (OpenCV) - Morphological operations
- PaddleOCR - Table + Layout detection
- YOLOv8 / CascadeTabNet - Advanced table detection (optional)
- python-docx - Word reconstruction with layout preservation
"""

import cv2
import numpy as np
from PIL import Image
from typing import List, Dict, Tuple, Optional, Union
from dataclasses import dataclass, field
import io

# OpenCV for classical CV
import cv2
import numpy as np

# PIL for image handling
from PIL import Image

# python-docx for Word export
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PaddleOCR for layout + table
try:
    from paddleocr import PaddleOCR, PPStructure
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False
    print("⚠️ PaddleOCR not installed")
except Exception as e:
    PADDLEOCR_AVAILABLE = False
    print(f"⚠️ PaddleOCR import error (torch DLL): {e}")

# PyMuPDF for PDF handling
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False


@dataclass
class BBox:
    """Bounding box với coordinates"""
    x1: int
    y1: int
    x2: int
    y2: int
    
    @property
    def width(self) -> int:
        return self.x2 - self.x1
    
    @property
    def height(self) -> int:
        return self.y2 - self.y1
    
    @property
    def center(self) -> Tuple[int, int]:
        return ((self.x1 + self.x2) // 2, (self.y1 + self.y2) // 2)
    
    def to_tuple(self) -> Tuple[int, int, int, int]:
        return (self.x1, self.y1, self.x2, self.y2)


@dataclass
class TextElement:
    """Một phần tử text với vị trí"""
    text: str
    bbox: BBox
    confidence: float = 1.0
    font_size: float = 12.0
    is_bold: bool = False
    element_type: str = "text"  # "text", "title", "header"


@dataclass
class TableCell:
    """Một cell trong bảng"""
    text: str
    bbox: BBox
    row: int
    col: int
    rowspan: int = 1
    colspan: int = 1
    confidence: float = 1.0


@dataclass
class Table:
    """Bảng với cells và structure"""
    bbox: BBox
    cells: List[TableCell] = field(default_factory=list)
    rows: int = 0
    cols: int = 0
    confidence: float = 1.0


@dataclass
class PageLayout:
    """Layout của một trang"""
    width: int
    height: int
    text_elements: List[TextElement] = field(default_factory=list)
    tables: List[Table] = field(default_factory=list)
    page_num: int = 1


def _crop_bbox(image: np.ndarray, bbox: BBox, padding: int = 2) -> np.ndarray:
    """Crop a bbox from the page image and keep bounds valid."""
    h, w = image.shape[:2]
    x1 = max(0, bbox.x1 - padding)
    y1 = max(0, bbox.y1 - padding)
    x2 = min(w, bbox.x2 + padding)
    y2 = min(h, bbox.y2 + padding)

    if x2 <= x1 or y2 <= y1:
        return image[0:0, 0:0]

    return image[y1:y2, x1:x2].copy()


def _bbox_overlap_ratio(a: BBox, b: BBox) -> float:
    """Return overlap ratio relative to bbox a."""
    x1 = max(a.x1, b.x1)
    y1 = max(a.y1, b.y1)
    x2 = min(a.x2, b.x2)
    y2 = min(a.y2, b.y2)

    if x2 <= x1 or y2 <= y1:
        return 0.0

    inter = (x2 - x1) * (y2 - y1)
    area_a = max(1, a.width * a.height)
    return inter / area_a


class ClassicalCVTableDetector:
    """
    Table detection sử dụng Classical Computer Vision (OpenCV)
    Morphological operations để detect lines và cells
    """
    
    def __init__(self, min_table_area: int = 1000):
        self.min_table_area = min_table_area
    
    def detect_tables(self, image: np.ndarray) -> List[Table]:
        """
        Detect tables sử dụng morphological operations
        
        Args:
            image: Grayscale image array
            
        Returns:
            List of Table objects
        """
        # Ensure grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()
        
        h, w = gray.shape
        
        # Binarize
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        
        # Detect horizontal lines
        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(20, w // 30), 1))
        horizontal = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel)
        
        # Detect vertical lines
        vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(20, h // 30)))
        vertical = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical_kernel)
        
        # Combine to get grid
        grid = cv2.add(horizontal, vertical)
        
        # Find table contours
        contours, _ = cv2.findContours(grid, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        tables = []
        for cnt in contours:
            x, y, ww, hh = cv2.boundingRect(cnt)
            area = ww * hh
            
            if area < self.min_table_area:
                continue
            
            bbox = BBox(x, y, x + ww, y + hh)
            
            # Extract cells from this table region
            cells, rows, cols = self._extract_cells(binary, bbox)
            
            if cells:
                table = Table(
                    bbox=bbox,
                    cells=cells,
                    rows=rows,
                    cols=cols,
                    confidence=0.85
                )
                tables.append(table)
        
        return tables
    
    def _extract_cells(self, binary: np.ndarray, table_bbox: BBox) -> Tuple[List[TableCell], int, int]:
        """Extract cells từ table region"""
        # Crop table region
        roi = binary[table_bbox.y1:table_bbox.y2, table_bbox.x1:table_bbox.x2]
        
        # Find row/column separators
        h_proj = cv2.reduce(roi, 1, cv2.REDUCE_AVG).flatten()
        v_proj = cv2.reduce(roi, 0, cv2.REDUCE_AVG).flatten()
        
        row_seps = np.where(h_proj > 20)[0]
        col_seps = np.where(v_proj > 20)[0]
        
        # Cluster separators
        row_bounds = self._cluster_separators(row_seps)
        col_bounds = self._cluster_separators(col_seps)
        
        if not row_bounds or not col_bounds:
            return [], 0, 0
        
        # Add edges
        rows = [0] + row_bounds + [roi.shape[0]]
        cols = [0] + col_bounds + [roi.shape[1]]
        
        # Create cells
        cells = []
        for ri in range(len(rows) - 1):
            for ci in range(len(cols) - 1):
                cell_bbox = BBox(
                    table_bbox.x1 + cols[ci],
                    table_bbox.y1 + rows[ri],
                    table_bbox.x1 + cols[ci + 1],
                    table_bbox.y1 + rows[ri + 1]
                )
                
                # Skip tiny cells (increase minimum size to avoid OCR errors)
                if cell_bbox.width < 30 or cell_bbox.height < 30:
                    continue
                
                cell = TableCell(
                    text="",  # Will be filled by OCR
                    bbox=cell_bbox,
                    row=ri,
                    col=ci
                )
                cells.append(cell)
        
        return cells, len(rows) - 1, len(cols) - 1
    
    def _cluster_separators(self, indices: np.ndarray, gap: int = 5) -> List[int]:
        """Cluster close indices"""
        if len(indices) == 0:
            return []
        
        groups = [[indices[0]]]
        for i in indices[1:]:
            if i - groups[-1][-1] <= gap:
                groups[-1].append(i)
            else:
                groups.append([i])
        
        return [int(np.mean(g)) for g in groups]


class PaddleOCRLayoutAnalyzer:
    """
    Layout analysis sử dụng PaddleOCR PPStructure
    Detect layout + table + text với coordinates
    """
    
    def __init__(self, lang: str = 'vi', use_gpu: bool = False):
        if not PADDLEOCR_AVAILABLE:
            raise ImportError("PaddleOCR not installed. Run: pip install paddleocr")
        
        # PPStructure for layout analysis
        # Minimal arguments for compatibility with PaddleOCR 3.3+
        self.engine = PPStructure(
            lang=lang,
            table=True,  # Enable table recognition
            ocr=True     # Enable OCR
        )
    
    def analyze_layout(self, image: Union[np.ndarray, Image.Image]) -> PageLayout:
        """
        Analyze page layout với PaddleOCR
        
        Returns:
            PageLayout với text elements và tables
        """
        # Convert to numpy if PIL
        if isinstance(image, Image.Image):
            img_array = np.array(image)
        else:
            img_array = image
        
        h, w = img_array.shape[:2]
        
        # Run PPStructure
        result = self.engine(img_array)
        
        layout = PageLayout(width=w, height=h)
        
        for item in result:
            region_type = item.get('type', 'text')
            bbox_list = item.get('bbox', [0, 0, w, h])
            
            bbox = BBox(
                int(bbox_list[0]),
                int(bbox_list[1]),
                int(bbox_list[2]),
                int(bbox_list[3])
            )
            
            if region_type == 'table':
                # Table detected
                table_result = item.get('res', {})
                table = self._parse_table_result(table_result, bbox)
                layout.tables.append(table)
            
            elif region_type in ['text', 'title', 'figure']:
                # Text element
                text_content = self._extract_text_from_result(item)
                
                if text_content:
                    element = TextElement(
                        text=text_content,
                        bbox=bbox,
                        confidence=item.get('score', 1.0),
                        element_type=region_type
                    )
                    layout.text_elements.append(element)
        
        return layout
    
    def _parse_table_result(self, table_result: Dict, bbox: BBox) -> Table:
        """Parse table result from PaddleOCR"""
        # PaddleOCR returns table HTML or cell structure
        cells = []
        
        # Try to extract cells
        if 'cells' in table_result:
            for cell_data in table_result['cells']:
                cell = TableCell(
                    text=cell_data.get('text', ''),
                    bbox=bbox,  # Approximate, would need cell bbox
                    row=cell_data.get('row', 0),
                    col=cell_data.get('col', 0),
                    rowspan=cell_data.get('rowspan', 1),
                    colspan=cell_data.get('colspan', 1)
                )
                cells.append(cell)
        
        max_row = max([c.row for c in cells], default=0) + 1
        max_col = max([c.col for c in cells], default=0) + 1
        
        return Table(
            bbox=bbox,
            cells=cells,
            rows=max_row,
            cols=max_col,
            confidence=table_result.get('score', 0.9)
        )
    
    def _extract_text_from_result(self, item: Dict) -> str:
        """Extract text từ OCR result"""
        if 'res' in item:
            res = item['res']
            if isinstance(res, list):
                texts = [r.get('text', '') for r in res if isinstance(r, dict)]
                return ' '.join(texts)
            elif isinstance(res, str):
                return res
        
        return item.get('text', '')


class LayoutToWordConverter:
    """
    Convert PageLayout to Word document giữ nguyên vị trí
    """
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        self.doc = Document()
    
    def add_page(self, layout: PageLayout, page_image: Optional[np.ndarray] = None, ocr_func=None):
        """
        Add một page layout vào Word document
        
        Args:
            layout: PageLayout object
            ocr_func: Function to OCR text elements/cells (optional)
        """
        # Sort elements by position (top to bottom)
        all_elements = []
        
        # Add text elements
        for elem in layout.text_elements:
            all_elements.append({
                'type': 'text',
                'y': elem.bbox.y1,
                'data': elem
            })
        
        # Add tables
        for table in layout.tables:
            all_elements.append({
                'type': 'table',
                'y': table.bbox.y1,
                'data': table
            })
        
        # Sort by vertical position
        all_elements.sort(key=lambda x: x['y'])
        
        # Add to document
        for elem in all_elements:
            if elem['type'] == 'text':
                self._add_text_element(elem['data'])
            elif elem['type'] == 'table':
                self._add_table(elem['data'], page_image=page_image, ocr_func=ocr_func)
    
    def _add_text_element(self, element: TextElement):
        """Add text element to document"""
        para = self.doc.add_paragraph()
        run = para.add_run(element.text)
        
        # Font formatting
        run.font.size = Pt(max(8, min(72, element.font_size)))
        run.bold = element.is_bold or element.element_type == 'title'
        
        # Paragraph alignment
        if element.element_type == 'title':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def _add_table(self, table: Table, page_image: Optional[np.ndarray] = None, ocr_func=None):
        """Add table to document"""
        if table.rows == 0 or table.cols == 0:
            return
        
        # Create Word table
        doc_table = self.doc.add_table(rows=table.rows, cols=table.cols)
        doc_table.style = 'Table Grid'
        doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Fill cells
        cell_map = {(c.row, c.col): c for c in table.cells}
        
        for row_idx in range(table.rows):
            for col_idx in range(table.cols):
                key = (row_idx, col_idx)
                
                if key in cell_map:
                    cell = cell_map[key]
                    doc_cell = doc_table.cell(row_idx, col_idx)
                    
                    # Get cell text - prioritize existing text, then OCR if empty
                    cell_text = cell.text if cell.text else ""
                    
                    # OCR cell if function provided and cell is empty
                    if ocr_func and page_image is not None and (not cell_text or cell_text.strip() == ""):
                        try:
                            cell_image = _crop_bbox(page_image, cell.bbox)
                            ocr_text = ocr_func(cell_image) if cell_image.size > 0 else ""
                            if ocr_text and ocr_text.strip():
                                cell_text = ocr_text
                                cell.text = ocr_text  # Update cell text
                        except Exception as e:
                            # If OCR fails, keep empty or existing text
                            pass
                    
                    # Set text (ensure not empty string with just spaces)
                    doc_cell.text = cell_text.strip() if cell_text else ""
                    
                    # Handle spans
                    if cell.rowspan > 1 or cell.colspan > 1:
                        try:
                            merge_cell = doc_table.cell(
                                row_idx,
                                col_idx
                            ).merge(doc_table.cell(
                                row_idx + cell.rowspan - 1,
                                col_idx + cell.colspan - 1
                            ))
                        except:
                            pass
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def save(self, output_path: str):
        """Save document"""
        self.doc.save(output_path)
    
    def get_bytes(self) -> bytes:
        """Get document as bytes"""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


def pdf_to_word_with_layout(
    pdf_input: Union[str, bytes, io.BytesIO],
    output_path: str,
    method: str = 'paddleocr',  # 'paddleocr', 'opencv', 'hybrid'
    ocr_func=None,
    dpi: int = 300
) -> Tuple[bool, str]:
    """
    Convert PDF to Word giữ nguyên layout
    
    Args:
        pdf_input: PDF file path or bytes
        output_path: Output Word file path
        method: 'paddleocr', 'opencv', or 'hybrid'
        ocr_func: Function to OCR text (receives bbox, returns text)
        dpi: Image resolution
        
    Returns:
        (success, message)
    """
    if not PYMUPDF_AVAILABLE:
        return False, "PyMuPDF not installed"
    
    # Load PDF
    if isinstance(pdf_input, str):
        doc = fitz.open(pdf_input)
    elif isinstance(pdf_input, bytes):
        doc = fitz.open(stream=pdf_input, filetype="pdf")
    elif isinstance(pdf_input, io.BytesIO):
        pdf_input.seek(0)
        doc = fitz.open(stream=pdf_input.read(), filetype="pdf")
    else:
        return False, f"Unsupported input type: {type(pdf_input)}"
    
    # Initialize converter
    converter = LayoutToWordConverter()
    
    # Initialize analyzer based on method
    if method in ('paddleocr', 'hybrid') and PADDLEOCR_AVAILABLE:
        analyzer = PaddleOCRLayoutAnalyzer(lang='vi')
        use_paddleocr = True
    else:
        analyzer = None
        use_paddleocr = False
    
    table_detector = ClassicalCVTableDetector()
    
    # Get total pages BEFORE processing (important: before closing doc)
    total_pages = len(doc)
    
    # Process each page
    for page_num in range(total_pages):
        page = doc.load_page(page_num)
        
        # Convert to image
        zoom = dpi / 72
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        
        img_data = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_data))
        img_array = np.array(image)
        
        h, w = img_array.shape[:2]
        layout = PageLayout(width=w, height=h, page_num=page_num + 1)

        # Detect tables first so overlapping digital text blocks can be skipped.
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY) if len(img_array.shape) == 3 else img_array
        layout.tables = table_detector.detect_tables(gray)

        # Prefer direct PDF text extraction when available because it creates a much more usable docx.
        text_dict = page.get_text("dict")
        scale = dpi / 72.0
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue

            block_bbox = BBox(
                int(block["bbox"][0] * scale),
                int(block["bbox"][1] * scale),
                int(block["bbox"][2] * scale),
                int(block["bbox"][3] * scale)
            )

            if any(_bbox_overlap_ratio(block_bbox, table.bbox) > 0.6 for table in layout.tables):
                continue

            lines = []
            font_sizes = []
            is_bold = False
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                line_text = "".join(span.get("text", "") for span in spans).strip()
                if not line_text:
                    continue
                lines.append(line_text)
                for span in spans:
                    if span.get("size"):
                        font_sizes.append(float(span["size"]))
                    if "bold" in str(span.get("font", "")).lower():
                        is_bold = True

            if not lines:
                continue

            font_size = max(font_sizes) if font_sizes else 11.0
            layout.text_elements.append(
                TextElement(
                    text="\n".join(lines),
                    bbox=block_bbox,
                    confidence=1.0,
                    font_size=font_size,
                    is_bold=is_bold,
                    element_type="title" if font_size >= 14 else "text"
                )
            )

        # If there is no embedded text, fall back to OCR.
        if not layout.text_elements:
            if use_paddleocr:
                try:
                    analyzed_layout = analyzer.analyze_layout(img_array)
                    if analyzed_layout.text_elements:
                        layout.text_elements = analyzed_layout.text_elements
                    if analyzed_layout.tables:
                        layout.tables = analyzed_layout.tables
                except Exception:
                    pass

            if not layout.text_elements and ocr_func:
                try:
                    ocr_text = (ocr_func(img_array) or "").strip()
                except Exception:
                    ocr_text = ""

                if ocr_text:
                    layout.text_elements.append(
                        TextElement(
                            text=ocr_text,
                            bbox=BBox(0, 0, w, h),
                            confidence=0.8,
                            font_size=11.0,
                            element_type="text"
                        )
                    )
        
        # Add page to Word
        converter.add_page(layout, page_image=img_array, ocr_func=ocr_func)
        
        # Add page break if not last page
        if page_num < total_pages - 1:
            converter.doc.add_page_break()
    
    # Close document AFTER processing
    doc.close()
    
    # Save
    try:
        converter.save(output_path)
        return True, f"Converted {total_pages} pages successfully"
    except Exception as e:
        return False, f"Error saving: {str(e)}"


# Export functions
__all__ = [
    'ClassicalCVTableDetector',
    'PaddleOCRLayoutAnalyzer',
    'LayoutToWordConverter',
    'PageLayout',
    'Table',
    'TableCell',
    'TextElement',
    'BBox',
    'pdf_to_word_with_layout',
    'PADDLEOCR_AVAILABLE',
    'DOCX_AVAILABLE',
    'PYMUPDF_AVAILABLE'
]
