"""
Benchmark a PDF -> Word conversion against an expected text file.

Usage:
    python tools/benchmark_pdf_to_word.py --pdf sample.pdf --expected sample.txt --use-ocr
"""

from __future__ import annotations

import argparse
import io
import sys
import unicodedata
from collections import Counter
from difflib import SequenceMatcher
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from docx import Document  # type: ignore
from pdf_to_word import convert_pdf_to_word


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.splitlines())
    text = "\n".join(line for line in text.splitlines() if line.strip())
    return " ".join(text.split())


def docx_to_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def ratio(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()


def build_ocr_callback():
    from vietnamese_ocr_advanced import VietnameseOCRAdvanced

    ocr = VietnameseOCRAdvanced(device="cpu", enable_all=True)

    def callback(image):
        result = ocr.recognize(
            image,
            engines=["paddleocr", "tesseract", "easyocr", "vietocr"],
            voting_method="best",
            preprocess="medium",
            fast_mode=True,
            verbose=False,
        )
        return result.text or ""

    return callback


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--pdf", required=True, help="Input PDF path")
    parser.add_argument("--expected", required=True, help="Expected text file path")
    parser.add_argument("--dpi", type=int, default=220)
    parser.add_argument("--no-tables", action="store_true")
    parser.add_argument("--use-ocr", action="store_true")
    args = parser.parse_args()

    pdf_path = Path(args.pdf)
    expected_path = Path(args.expected)
    expected_text = expected_path.read_text(encoding="utf-8")

    ocr_callback = build_ocr_callback() if args.use_ocr else None
    result = convert_pdf_to_word(
        pdf_input=pdf_path,
        dpi=args.dpi,
        extract_tables=not args.no_tables,
        ocr_callback=ocr_callback,
    )

    if not result.success:
        print(f"Conversion failed: {result.error_message}")
        return 1

    extracted_text = normalize_text(result.extracted_text)
    expected_text_norm = normalize_text(expected_text)
    docx_text = normalize_text(docx_to_text(result.docx_bytes))

    print(f"Pages converted: {result.pages_converted}")
    print(f"OCR pages: {result.ocr_pages}")
    print(f"Tables: {result.table_count}")
    print(f"Method: {result.method_used}")
    print()
    print(f"Extracted text similarity: {ratio(extracted_text, expected_text_norm):.4f}")
    print(f"DOCX text similarity: {ratio(docx_text, expected_text_norm):.4f}")
    print()

    strategies = Counter(page.strategy for page in result.pages)
    print("Page strategies:")
    for name, count in sorted(strategies.items()):
        print(f"  {name}: {count}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
